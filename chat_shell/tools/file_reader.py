"""
File reader tool for parsing various file formats.
"""

import io
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from pydantic import BaseModel, Field

from .base import BaseTool, ToolInput, ToolOutput

logger = logging.getLogger(__name__)


class FileReaderInput(ToolInput):
    """Input schema for file reader tool."""
    file_path: str = Field(..., description="Path to the file to read")
    max_lines: Optional[int] = Field(default=None, description="Maximum number of lines to read (for text files)")
    extract_metadata: bool = Field(default=True, description="Extract file metadata")


class FileMetadata(BaseModel):
    """File metadata."""
    file_name: str
    file_size: int
    file_type: str
    encoding: Optional[str] = None
    line_count: Optional[int] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None


class FileReaderOutput(ToolOutput):
    """Output schema for file reader tool."""
    content: str = ""
    metadata: Optional[FileMetadata] = None
    sections: List[Dict[str, Any]] = Field(default_factory=list)


class FileReaderTool(BaseTool):
    """File reader tool supporting multiple formats.

    Supports:
    - Text files (.txt, .md, .py, .js, etc.)
    - PDF files (.pdf) - requires PyPDF2 or pdfplumber
    - Word documents (.docx) - requires python-docx
    - Excel files (.xlsx, .csv) - requires openpyxl, pandas
    - JSON files (.json)

    Example:
        tool = FileReaderTool()
        result = await tool.execute(FileReaderInput(file_path="document.pdf"))
    """

    name = "file_reader"
    description = (
        "Read and extract content from files. "
        "Supports PDF, Word, Excel, CSV, JSON, Markdown, and text files. "
        "Returns file content as text with metadata."
    )
    input_schema = FileReaderInput

    # Supported file types
    SUPPORTED_EXTENSIONS = {
        '.txt', '.md', '.markdown', '.rst',
        '.py', '.js', '.ts', '.jsx', '.tsx',
        '.java', '.c', '.cpp', '.h', '.hpp',
        '.go', '.rs', '.rb', '.php', '.swift',
        '.kt', '.scala', '.r', '.m', '.mm',
        '.html', '.htm', '.css', '.scss', '.sass',
        '.json', '.yaml', '.yml', '.toml', '.xml',
        '.csv', '.tsv',
        '.pdf', '.docx', '.xlsx',
    }

    async def execute(self, input_data: FileReaderInput) -> ToolOutput:
        """Execute file reader."""
        file_path = Path(input_data.file_path)

        # Validate path
        if not file_path.exists():
            return ToolOutput(
                result="",
                error=f"File not found: {file_path}"
            )

        if not file_path.is_file():
            return ToolOutput(
                result="",
                error=f"Path is not a file: {file_path}"
            )

        # Check extension
        extension = file_path.suffix.lower()
        if extension not in self.SUPPORTED_EXTENSIONS:
            return ToolOutput(
                result="",
                error=f"Unsupported file type: {extension}. "
                      f"Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        try:
            # Route to appropriate reader
            if extension == '.pdf':
                return await self._read_pdf(file_path, input_data)
            elif extension == '.docx':
                return await self._read_docx(file_path, input_data)
            elif extension == '.xlsx':
                return await self._read_xlsx(file_path, input_data)
            elif extension == '.csv':
                return await self._read_csv(file_path, input_data)
            elif extension == '.json':
                return await self._read_json(file_path, input_data)
            else:
                # Text-based files
                return await self._read_text(file_path, input_data)

        except Exception as e:
            logger.error(f"Failed to read file {file_path}: {e}")
            return ToolOutput(
                result="",
                error=f"Failed to read file: {str(e)}"
            )

    async def _read_text(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read text-based files."""
        # Try to detect encoding if chardet is available
        try:
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Sample first 10KB
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8') or 'utf-8'
        except ImportError:
            encoding = 'utf-8'

        # Read file
        content = file_path.read_text(encoding=encoding, errors='replace')

        # Apply line limit if specified
        if input_data.max_lines:
            lines = content.split('\n')[:input_data.max_lines]
            content = '\n'.join(lines)
            if len(content.split('\n')) == input_data.max_lines:
                content += "\n\n[Content truncated...]"

        # Calculate metadata
        metadata = None
        if input_data.extract_metadata:
            lines = content.split('\n')
            words = content.split()
            metadata = FileMetadata(
                file_name=file_path.name,
                file_size=file_path.stat().st_size,
                file_type=f"text/{file_path.suffix.lstrip('.')}",
                encoding=encoding,
                line_count=len(lines),
                word_count=len(words)
            )

        return FileReaderOutput(
            result=content,
            content=content,
            metadata=metadata
        )

    async def _read_pdf(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read PDF files."""
        try:
            import PyPDF2
        except ImportError:
            return ToolOutput(
                result="",
                error="PyPDF2 is required for PDF reading. Install with: pip install PyPDF2"
            )

        content_parts = []
        metadata = None

        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)

            for i, page in enumerate(reader.pages):
                if input_data.max_lines and i >= input_data.max_lines:
                    content_parts.append("\n[Content truncated...]")
                    break

                page_text = page.extract_text()
                if page_text:
                    content_parts.append(f"--- Page {i + 1} ---\n{page_text}")

            if input_data.extract_metadata:
                metadata = FileMetadata(
                    file_name=file_path.name,
                    file_size=file_path.stat().st_size,
                    file_type="application/pdf",
                    page_count=num_pages
                )

        content = "\n\n".join(content_parts)
        return FileReaderOutput(
            result=content,
            content=content,
            metadata=metadata
        )

    async def _read_docx(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read Word documents."""
        try:
            import docx
        except ImportError:
            return ToolOutput(
                result="",
                error="python-docx is required for Word reading. Install with: pip install python-docx"
            )

        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        if input_data.max_lines:
            paragraphs = paragraphs[:input_data.max_lines]
            if len(doc.paragraphs) > input_data.max_lines:
                paragraphs.append("[Content truncated...]")

        content = "\n\n".join(paragraphs)

        metadata = None
        if input_data.extract_metadata:
            metadata = FileMetadata(
                file_name=file_path.name,
                file_size=file_path.stat().st_size,
                file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                word_count=len(content.split())
            )

        return FileReaderOutput(
            result=content,
            content=content,
            metadata=metadata
        )

    async def _read_xlsx(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read Excel files."""
        try:
            import pandas as pd
        except ImportError:
            return ToolOutput(
                result="",
                error="pandas is required for Excel reading. Install with: pip install pandas openpyxl"
            )

        try:
            # Read all sheets
            xls = pd.ExcelFile(file_path)
            content_parts = []

            for sheet_name in xls.sheet_names[:3]:  # Limit to first 3 sheets
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                if input_data.max_lines:
                    df = df.head(input_data.max_lines)

                content_parts.append(f"=== Sheet: {sheet_name} ===")
                content_parts.append(df.to_string(index=False))
                content_parts.append("")

            if len(xls.sheet_names) > 3:
                content_parts.append(f"[{len(xls.sheet_names) - 3} more sheets...]")

            content = "\n".join(content_parts)

            metadata = None
            if input_data.extract_metadata:
                metadata = FileMetadata(
                    file_name=file_path.name,
                    file_size=file_path.stat().st_size,
                    file_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

            return FileReaderOutput(
                result=content,
                content=content,
                metadata=metadata
            )

        except Exception as e:
            return ToolOutput(
                result="",
                error=f"Failed to read Excel file: {str(e)}"
            )

    async def _read_csv(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read CSV files."""
        try:
            import pandas as pd
        except ImportError:
            # Fallback to standard library
            import csv

            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)

            if input_data.max_lines:
                rows = rows[:input_data.max_lines]

            # Format as table
            content_lines = []
            for row in rows:
                content_lines.append(" | ".join(row))

            content = "\n".join(content_lines)

            return FileReaderOutput(
                result=content,
                content=content
            )

        # Use pandas for better formatting
        df = pd.read_csv(file_path)

        if input_data.max_lines:
            df = df.head(input_data.max_lines)

        content = df.to_string(index=False)

        metadata = None
        if input_data.extract_metadata:
            metadata = FileMetadata(
                file_name=file_path.name,
                file_size=file_path.stat().st_size,
                file_type="text/csv",
                line_count=len(df)
            )

        return FileReaderOutput(
            result=content,
            content=content,
            metadata=metadata
        )

    async def _read_json(self, file_path: Path, input_data: FileReaderInput) -> ToolOutput:
        """Read JSON files."""
        import json

        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        # Pretty print JSON
        content = json.dumps(data, indent=2, ensure_ascii=False)

        metadata = None
        if input_data.extract_metadata:
            metadata = FileMetadata(
                file_name=file_path.name,
                file_size=file_path.stat().st_size,
                file_type="application/json"
            )

        return FileReaderOutput(
            result=content,
            content=content,
            metadata=metadata
        )
